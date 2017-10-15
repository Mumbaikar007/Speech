from docx import Document
import speechtotext
from docx.shared import Inches

def word_main():
    document = Document()
    i = speechtotext.stot()
    print(i)

    #document.add_heading("python doc")
    document.add_paragraph(i)


    document.save('Demo.docx') 
