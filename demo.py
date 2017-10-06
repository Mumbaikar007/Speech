from docx import Document
from docx.shared import Inches

document = Document()


document.add_heading("python doc")
document.add_paragraph("This is a demo")


document.save('Demo.docx')